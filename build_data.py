# -*- coding: utf-8 -*-
import docx
import re
import json

def z2h(s):
    return s.translate(str.maketrans('０１２３４５６７８９', '0123456789'))

# ---------- Japanese ----------
def parse_ja(fname):
    d = docx.Document(fname)
    paras = [p.text for p in d.paragraphs]

    qa_lines = []
    footnote_lines = []
    in_footnotes = False
    for line in paras:
        t = line.strip()
        if not t:
            continue
        if re.match(r'^\(\d+\)$', t):
            in_footnotes = True
        if in_footnotes:
            footnote_lines.append(t)
        else:
            qa_lines.append(t)

    # footnote dict: pairs of "(N)" then citation text
    footnotes = {}
    i = 0
    while i < len(footnote_lines) - 1:
        m = re.match(r'^\((\d+)\)$', footnote_lines[i])
        if m:
            footnotes[int(m.group(1))] = footnote_lines[i + 1]
            i += 2
        else:
            i += 1

    entries = {}
    i = 0
    while i < len(qa_lines):
        m = re.match(r'^問([0-9０-９]+)　(.*)$', qa_lines[i])
        if m:
            qid = int(z2h(m.group(1)))
            qtext = m.group(2)
            atext = qa_lines[i + 1] if i + 1 < len(qa_lines) else ''
            i += 2

            combined = qtext + ' ' + atext
            ref_nums = [int(n) for n in re.findall(r'\((\d+)\)', combined)]
            refs = [footnotes[n] for n in ref_nums if n in footnotes]

            qtext_clean = re.sub(r'\(\d+\)', '', qtext).strip()
            atext_clean = re.sub(r'\(\d+\)', '', atext).strip()

            entries[qid] = {
                'id': qid,
                'question': qtext_clean,
                'answer': atext_clean,
                'refs': refs,
            }
        else:
            i += 1
    return entries

# ---------- English ----------
def parse_en(fname):
    d = docx.Document(fname)
    paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]

    entries = {}
    i = 0
    cur = None
    while i < len(paras):
        t = paras[i]
        mq = re.match(r'^Q\.\s*(\d+)\.\s*(.*)$', t)
        ma = re.match(r'^A\.\s*(.*)$', t)
        if mq:
            cur = {'id': int(mq.group(1)), 'questionEn': mq.group(2).strip(), 'answerEn': '', 'refsEn': []}
            entries[cur['id']] = cur
            i += 1
            continue
        if ma and cur is not None:
            cur['answerEn'] = ma.group(1).strip()
            i += 1
            continue
        if t == 'Back to Top':
            i += 1
            continue
        if cur is not None:
            # either an answer continuation or a scripture reference line
            if re.search(r'\d+:\d+', t):
                cur['refsEn'].append(t.rstrip('.'))
            else:
                cur['answerEn'] = (cur['answerEn'] + ' ' + t).strip()
        i += 1
    return entries

# ---------- Bible reference linking ----------
# YouVersion (bible.com) deep links: chapter/verse pages work without login.
#   Japanese - Shinkyodoyaku (新共同訳), version id 1819
#   English  - NIV, version id 111
JA_ALIASES = {
    'Ⅰコリント': '1CO', 'Ⅱコリント': '2CO', 'コリント': '1CO',
    'Ⅰテモテ': '1TI', 'Ⅱテモテ': '2TI',
    'Ⅰペテロ': '1PE', 'Ⅱペテロ': '2PE',
    'Ⅰヨハネ': '1JN', 'Ⅱヨハネ': '2JN', 'Ⅲヨハネ': '3JN',
    'Ⅰ列王': '1KI', 'Ⅱ列王': '2KI',
    'Ⅰ歴代': '1CH', 'Ⅱ歴代': '2CH',
    'Ⅰサムエル': '1SA', 'Ⅱサムエル': '2SA',
    'Ⅰテサロニケ': '1TH', 'Ⅱテサロニケ': '2TH',
    '創世': 'GEN', '出エジプト': 'EXO', 'レビ': 'LEV', '民数': 'NUM', '申命': 'DEU',
    'ヨシュア': 'JOS', '士師': 'JDG', 'ルツ': 'RUT', 'エズラ': 'EZR', 'ネヘミヤ': 'NEH',
    'エステル': 'EST', 'ヨブ': 'JOB', '詩篇': 'PSA', '詩': 'PSA', '箴言': 'PRO',
    '伝道': 'ECC', '雅歌': 'SNG', 'イザヤ': 'ISA', 'エレミヤ': 'JER', '哀歌': 'LAM',
    'エゼキエル': 'EZK', 'ダニエル': 'DAN', 'ホセア': 'HOS', 'ヨエル': 'JOL',
    'アモス': 'AMO', 'オバデヤ': 'OBA', 'ヨナ': 'JON', 'ミカ': 'MIC', 'ナホム': 'NAM',
    'ハバクク': 'HAB', 'ゼパニヤ': 'ZEP', 'ハガイ': 'HAG', 'ゼカリヤ': 'ZEC', 'マラキ': 'MAL',
    'マタイ': 'MAT', 'マルコ': 'MRK', 'ルカ': 'LUK', 'ヨハネ': 'JHN', '使徒': 'ACT',
    'ローマ': 'ROM', 'ガラテヤ': 'GAL', 'エペソ': 'EPH', 'ピリピ': 'PHP', 'コロサイ': 'COL',
    'テトス': 'TIT', 'ピレモン': 'PHM', 'ヘブル': 'HEB', 'ヤコブ': 'JAS', 'ユダ': 'JUD',
}

EN_ALIASES_RAW = {
    'Gen': 'GEN', 'Ex': 'EXO', 'Lev': 'LEV', 'Num': 'NUM', 'Deut': 'DEU',
    'Josh': 'JOS', 'Judg': 'JDG', 'Ruth': 'RUT', '1 Sam': '1SA', '2 Sam': '2SA',
    '1 Kings': '1KI', '2 Kings': '2KI', '1 Chr': '1CH', '2 Chr': '2CH',
    'Ezra': 'EZR', 'Neh': 'NEH', 'Esther': 'EST', 'Job': 'JOB', 'Ps': 'PSA',
    'Prov': 'PRO', 'Ecc': 'ECC', 'Eccl': 'ECC', 'Song': 'SNG', 'Isa': 'ISA',
    'Jer': 'JER', 'Lam': 'LAM', 'Ezek': 'EZK', 'Dan': 'DAN', 'Hos': 'HOS',
    'Joel': 'JOL', 'Amos': 'AMO', 'Obad': 'OBA', 'Jonah': 'JON', 'Mic': 'MIC',
    'Nah': 'NAM', 'Hab': 'HAB', 'Zeph': 'ZEP', 'Hag': 'HAG', 'Zech': 'ZEC',
    'Mal': 'MAL', 'Matt': 'MAT', 'Mark': 'MRK', 'Luke': 'LUK', 'John': 'JHN',
    'Acts': 'ACT', 'Rom': 'ROM', '1 Cor': '1CO', '2 Cor': '2CO', 'Gal': 'GAL',
    'Eph': 'EPH', 'Phil': 'PHP', 'Col': 'COL', '1 Thess': '1TH', '2 Thess': '2TH',
    '1 Tim': '1TI', '2 Tim': '2TI', 'Titus': 'TIT', 'Philem': 'PHM', 'Heb': 'HEB',
    'Jas': 'JAS', '1 Pet': '1PE', '2 Pet': '2PE', '1 John': '1JN', '2 John': '2JN',
    '3 John': '3JN', 'Jude': 'JUD', 'Rev': 'REV',
}

def build_ja_alias_pattern():
    aliases = sorted(JA_ALIASES.keys(), key=len, reverse=True)
    return re.compile('(' + '|'.join(re.escape(a) for a in aliases) + ')')

def build_en_alias_pattern():
    aliases = sorted(EN_ALIASES_RAW.keys(), key=len, reverse=True)
    # optional trailing period, must be followed by whitespace (avoids matching inside other words)
    return re.compile('(' + '|'.join(re.escape(a) for a in aliases) + r')\.?(?=\s)')

JA_ALIAS_PATTERN = build_ja_alias_pattern()
EN_ALIAS_PATTERN = build_en_alias_pattern()

# Within a book's "run" of text (from the alias to the next alias / end of
# string) a token is either "chapter:verse[-verse]" (which also resets the
# current chapter, e.g. "Rev. 4:8; 15:4" -> two chapters of Revelation) or a
# bare "verse[-verse]" continuing the current chapter (e.g. "25:41, 46"), or
# -- if no chapter has been seen yet in this run -- a bare whole-chapter
# reference (e.g. "Ex. 12").
TOKEN_PATTERN = re.compile(r'(\d+):(\d+)(?:-(\d+))?|(\d+)(?:-(\d+))?')

def linkify(text, aliases, alias_pattern, url_template):
    alias_matches = list(alias_pattern.finditer(text))
    if not alias_matches:
        return [{'text': text}]

    segments = []
    cursor = 0
    for i, am in enumerate(alias_matches):
        a_start, a_end = am.start(), am.end()
        osis = aliases[am.group(1)]
        run_end = alias_matches[i + 1].start() if i + 1 < len(alias_matches) else len(text)

        if a_start > cursor:
            segments.append({'text': text[cursor:a_start]})

        current_chapter = None
        first_token = True
        pos = a_end
        for tm in TOKEN_PATTERN.finditer(text, a_end, run_end):
            if tm.start() > pos:
                segments.append({'text': text[pos:tm.start()]})
            if tm.group(1) is not None:
                chapter, v1, v2 = tm.group(1), tm.group(2), tm.group(3)
                current_chapter = chapter
                verse = f'{v1}-{v2}' if v2 else v1
                url = url_template.format(osis=osis, chapter=chapter, verse='.' + verse)
            else:
                n1, n2 = tm.group(4), tm.group(5)
                if current_chapter is None:
                    current_chapter = n1
                    url = url_template.format(osis=osis, chapter=n1, verse='')
                else:
                    verse = f'{n1}-{n2}' if n2 else n1
                    url = url_template.format(osis=osis, chapter=current_chapter, verse='.' + verse)
            label_start = a_start if first_token else tm.start()
            segments.append({'text': text[label_start:tm.end()], 'url': url})
            pos = tm.end()
            first_token = False

        if first_token:
            # no numeric token at all in this run - show the alias as plain text
            segments.append({'text': text[a_start:a_end]})
            pos = a_end
        elif pos < run_end:
            segments.append({'text': text[pos:run_end]})
        cursor = run_end

    if cursor < len(text):
        segments.append({'text': text[cursor:]})
    return segments

def linkify_ja(text):
    return linkify(text, JA_ALIASES, JA_ALIAS_PATTERN, 'https://www.bible.com/bible/1819/{osis}.{chapter}{verse}')

def linkify_en(text):
    return linkify(text, EN_ALIASES_RAW, EN_ALIAS_PATTERN, 'https://www.bible.com/bible/111/{osis}.{chapter}{verse}.NIV')

ja = parse_ja('ウェストミンスター小教理問答.docx')
en = parse_en('Westminster shorter Catechism.docx')

ids = sorted(set(ja.keys()) | set(en.keys()))
merged = []
unmatched_ja = []
unmatched_en = []
for qid in ids:
    j = ja.get(qid, {})
    e = en.get(qid, {})
    refs = j.get('refs', [])
    refs_en = e.get('refsEn', [])
    ref_links = [linkify_ja(r) for r in refs]
    ref_links_en = [linkify_en(r) for r in refs_en]
    for segs in ref_links:
        if not any('url' in s for s in segs):
            unmatched_ja.append((qid, segs))
    for segs in ref_links_en:
        if not any('url' in s for s in segs):
            unmatched_en.append((qid, segs))
    merged.append({
        'id': qid,
        'question': j.get('question', ''),
        'answer': j.get('answer', ''),
        'refs': ref_links,
        'questionEn': e.get('questionEn', ''),
        'answerEn': e.get('answerEn', ''),
        'refsEn': ref_links_en,
    })

if unmatched_ja:
    print('JA refs with no recognized citation:')
    for qid, segs in unmatched_ja:
        print(' ', qid, segs)
if unmatched_en:
    print('EN refs with no recognized citation:')
    for qid, segs in unmatched_en:
        print(' ', qid, segs)

print('total entries:', len(merged))
missing_ja = [m['id'] for m in merged if not m['question'] or not m['answer']]
missing_en = [m['id'] for m in merged if not m['questionEn'] or not m['answerEn']]
print('missing ja:', missing_ja)
print('missing en:', missing_en)

with open('catechism-data.json', 'w', encoding='utf-8') as f:
    json.dump(merged, f, ensure_ascii=False, indent=2)

# U+2028/U+2029 are valid inside JSON strings but are also treated as line
# terminators by (older) JS parsers, so embedding raw JSON straight into a
# <script> tag can throw a SyntaxError. Escape them defensively.
json_text = json.dumps(merged, ensure_ascii=False, indent=2)
json_text = json_text.replace(chr(0x2028), '\u2028').replace(chr(0x2029), '\u2029')

with open('catechism-data.js', 'w', encoding='utf-8') as f:
    f.write('const CATECHISM_DATA = ')
    f.write(json_text)
    f.write(';\n')

print('wrote catechism-data.json and catechism-data.js')
